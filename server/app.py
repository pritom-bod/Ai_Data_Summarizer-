from flask import Flask, request, send_file
import requests
from bs4 import BeautifulSoup
from docx import Document
import tempfile
import os

app = Flask(__name__)

@app.route("/scrape", methods=["POST"])
def scrape():
    try:
        data = request.get_json()
        urls = data.get("urls", [])
        if not urls:
            return "No URLs provided", 400

        doc = Document()
        doc.add_heading("Scraped Website Data", 0)

        for url in urls:
            try:
                response = requests.get(url, timeout=10)
                soup = BeautifulSoup(response.text, "html.parser")

                text = soup.get_text(separator="\n", strip=True)
                doc.add_page_break()
                doc.add_heading(url, level=2)
                doc.add_paragraph(text[:10000])  # Prevent too much

            except Exception as e:
                print(f"Failed to fetch {url}: {e}")

        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        doc.save(temp_path)

        return send_file(temp_path, as_attachment=True, download_name="website_data.docx")

    except Exception as e:
        print("Scraping error:", e)
        return "Server Error", 500

if __name__ == "__main__":
    app.run(debug=True)
